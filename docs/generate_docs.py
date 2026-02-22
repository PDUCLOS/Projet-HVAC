# -*- coding: utf-8 -*-
"""
Word documentation generator for the HVAC Market Analysis project.
===================================================================

Generates two comprehensive documents:
1. documentation_technique.docx — Full technical architecture & methodology
2. guide_utilisation.docx — Installation, configuration, accounts & deployment

Usage:
    python docs/generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def set_doc_styles(doc):
    """Configure document styles (Calibri 11, blue headings)."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 5):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0, 51, 102)


def add_cover_page(doc, title, subtitle):
    """Add a formatted cover page with project info."""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HVAC Market Analysis Project\nMetropolitan France")
    run.font.size = Pt(14)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Patrice DUCLOS\nSenior Data Analyst\n"
        "Data Science Lead Program — Jedha Bootcamp (Master's level)\n"
        "February 2026"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def add_table(doc, headers, rows):
    """Add a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)

    doc.add_paragraph("")
    return table


def add_bullets(doc, items):
    """Add a bulleted list."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    """Add a numbered list."""
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}")


# ================================================================
# DOCUMENT 1 : Technical Documentation
# ================================================================

def generate_technical_doc():
    """Generate the complete technical documentation."""
    doc = Document()
    set_doc_styles(doc)
    add_cover_page(
        doc,
        "Technical Documentation",
        "Architecture, Data Pipeline and ML Methodology"
    )

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Project Overview",
        "2. Technical Architecture",
        "3. Data Sources",
        "4. Collection Pipeline",
        "5. Data Processing",
        "6. Feature Engineering",
        "7. Database Schema",
        "8. ML Modeling",
        "9. Ridge Auto-correlation Analysis",
        "10. REST API (FastAPI)",
        "11. Streamlit Dashboard",
        "12. Deployment (Docker / Kubernetes)",
        "13. Orchestration (Airflow)",
        "14. Reinforcement Learning (demo)",
        "15. Unit Tests",
        "16. Training Modules Coverage",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # =================================================================
    # 1. PROJECT OVERVIEW
    # =================================================================
    doc.add_heading("1. Project Overview", level=1)

    doc.add_heading("1.1 Objective", level=2)
    doc.add_paragraph(
        "This project builds a complete data pipeline to analyze and predict "
        "HVAC equipment installations (heat pumps, air conditioning) "
        "across all 96 departments of Metropolitan France."
    )
    doc.add_paragraph(
        "This is a Data Science Lead portfolio project (Master's level) demonstrating "
        "mastery of the complete chain: open data collection, "
        "cleaning, star schema modeling, feature engineering, ML modeling, "
        "REST API, interactive dashboard, Docker/Kubernetes deployment and "
        "Airflow orchestration."
    )

    doc.add_heading("1.2 Proxy Strategy", level=2)
    doc.add_paragraph(
        "No direct HVAC sales data is available as open data. "
        "The project uses Energy Performance Diagnostics (DPE) "
        "from ADEME as a proxy: each DPE mentions the type of heating "
        "and air conditioning installed. By filtering DPEs mentioning a heat "
        "pump (PAC) or air conditioning, we obtain a reliable indicator "
        "of installation volume."
    )

    doc.add_heading("1.3 Scope", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Scope", "Metropolitan France (96 departments)"],
            ["Period", "July 2021 — February 2026 (~55 months)"],
            ["Granularity", "Month x Department"],
            ["DPE Volume", "Several million diagnostics"],
            ["ML Features", "~90 engineered features"],
            ["API", "6 REST endpoints (FastAPI)"],
            ["Dashboard", "6 interactive pages (Streamlit)"],
            ["Tests", "119 unit tests (pytest)"],
        ]
    )

    doc.add_heading("1.4 Key Results", level=2)
    add_table(doc,
        ["Model", "Val RMSE", "Val R2", "Test RMSE", "Test R2"],
        [
            ["Ridge (with lags)", "1.09", "0.9975", "0.96", "0.9982"],
            ["Ridge (exogenous)", "—", "—", "—", "—"],
            ["LightGBM", "5.45", "0.937", "8.37", "0.865"],
            ["Prophet", "13.75", "0.599", "12.05", "0.719"],
        ]
    )
    doc.add_paragraph(
        "Note: The very high Ridge R2 is due to auto-correlation "
        "via lag features. The 'Ridge exogenous' model evaluates the true "
        "contribution of weather and economic features (see section 9)."
    )
    doc.add_page_break()

    # =================================================================
    # 2. TECHNICAL ARCHITECTURE
    # =================================================================
    doc.add_heading("2. Technical Architecture", level=1)

    doc.add_heading("2.1 Project Structure", level=2)
    add_table(doc,
        ["Layer", "Directory", "Responsibility"],
        [
            ["Configuration", "config/settings.py", "Centralized parameters (frozen dataclasses)"],
            ["Collection", "src/collectors/", "5 plugin collectors (auto-registration)"],
            ["Processing", "src/processing/", "Cleaning, merging, feature engineering, outliers"],
            ["Database", "src/database/", "Star schema, CSV import → SQLite/PostgreSQL"],
            ["Analysis", "src/analysis/", "Automated EDA, correlations"],
            ["Modeling", "src/models/", "Ridge, LightGBM, Prophet, LSTM, RL demo"],
            ["REST API", "api/", "FastAPI — 6 prediction/monitoring endpoints"],
            ["Dashboard", "app/", "Streamlit — 6 interactive pages"],
            ["Orchestration", "airflow/dags/", "Airflow DAG complete pipeline"],
            ["Deployment", "kubernetes/, Dockerfile", "Docker multi-stage, K8s, Render"],
            ["Tests", "tests/", "119 unit tests (pytest)"],
        ]
    )

    doc.add_heading("2.2 Architecture Patterns", level=2)

    doc.add_heading("Plugin Registry (Collectors)", level=3)
    doc.add_paragraph(
        "Collectors use a plugin system based on __init_subclass__. "
        "Each concrete subclass of BaseCollector is automatically registered "
        "in the CollectorRegistry. To add a new data source, "
        "simply create a file in src/collectors/ with a class inheriting "
        "from BaseCollector — no manual configuration is required."
    )

    doc.add_heading("Frozen Dataclasses", level=3)
    doc.add_paragraph(
        "All configuration is defined via frozen (immutable) dataclasses. "
        "Values are loaded from the .env file via python-dotenv."
    )

    doc.add_heading("Idempotent Pipeline", level=3)
    doc.add_paragraph(
        "Each pipeline step can be re-run safely. Output files "
        "are overwritten. Duplicates are handled upstream."
    )

    doc.add_heading("2.3 Technologies", level=2)
    add_table(doc,
        ["Category", "Technologies"],
        [
            ["Language", "Python 3.10+"],
            ["Data", "pandas, numpy, SQLAlchemy"],
            ["ML", "scikit-learn (Ridge, StandardScaler, TimeSeriesSplit)"],
            ["Gradient Boosting", "LightGBM"],
            ["Time Series", "Prophet (Meta)"],
            ["Deep Learning", "PyTorch (LSTM, exploratory)"],
            ["Reinforcement Learning", "Gymnasium, Q-Learning"],
            ["Interpretability", "SHAP"],
            ["Visualization", "matplotlib, seaborn, plotly"],
            ["API", "FastAPI, Pydantic v2, uvicorn"],
            ["Dashboard", "Streamlit"],
            ["Database", "SQLite (default), PostgreSQL, SQL Server"],
            ["Containerization", "Docker (multi-stage), Docker Compose"],
            ["K8s Orchestration", "Kubernetes (Deployment, Ingress, CronJob)"],
            ["CI/CD", "Render (PaaS), deploy.sh"],
            ["Tests", "pytest"],
            ["HTTP", "requests (automatic retry)"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 3. DATA SOURCES
    # =================================================================
    doc.add_heading("3. Data Sources", level=1)
    doc.add_paragraph(
        "All sources are Open Data — no API key is required."
    )

    doc.add_heading("3.1 DPE ADEME (target variable)", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Source", "data.ademe.fr (paginated JSON API)"],
            ["Scope", "96 departments of Metropolitan France"],
            ["Period", "July 2021+ (DPE v2)"],
            ["Granularity", "1 row per diagnostic"],
            ["Heat pump detection", "Regex on heating/cooling: PAC, heat pump, thermodynamic"],
            ["Collector", "DpeCollector (src/collectors/dpe.py)"],
        ]
    )

    doc.add_heading("3.2 Open-Meteo (weather features)", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Source", "archive-api.open-meteo.com"],
            ["Scope", "96 reference cities (prefectures)"],
            ["Variables", "temperature max/min/mean, precipitation, wind"],
            ["Derived", "HDD (Heating Degree Days), CDD (Cooling Degree Days)"],
            ["Collector", "WeatherCollector (src/collectors/weather.py)"],
        ]
    )

    doc.add_heading("3.3 INSEE BDM (economic features)", level=2)
    add_table(doc,
        ["Series", "idbank", "Description"],
        [
            ["confiance_menages", "001759970", "Synthetic confidence indicator"],
            ["climat_affaires_industrie", "001565530", "Business climate all sectors"],
            ["climat_affaires_batiment", "001586808", "Business climate construction"],
            ["opinion_achats_importants", "001759974", "Major purchases opportunity"],
            ["situation_financiere_future", "001759972", "Future financial situation"],
            ["ipi_industrie_manuf", "010768261", "Manufacturing industry IPI"],
        ]
    )

    doc.add_heading("3.4 Eurostat (IPI HVAC)", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Dataset", "sts_inpr_m (Industrial Production)"],
            ["NACE Codes", "C28 (machinery), C2825 (AC equipment)"],
            ["Geo filter", "France (FR)"],
            ["Collector", "EurostatCollector (src/collectors/eurostat_col.py)"],
        ]
    )

    doc.add_heading("3.5 SITADEL (building permits)", level=2)
    doc.add_paragraph(
        "Source currently being migrated to the SDES DiDo API. "
        "The SitadelCollector exists but is not fully functional."
    )
    doc.add_page_break()

    # =================================================================
    # 4. COLLECTION PIPELINE
    # =================================================================
    doc.add_heading("4. Collection Pipeline", level=1)

    doc.add_heading("4.1 BaseCollector Architecture", level=2)
    doc.add_paragraph("All collectors inherit from BaseCollector:")
    add_bullets(doc, [
        "HTTP session with automatic retry (codes 429, 500-504)",
        "Helpers fetch_json(), fetch_xml(), fetch_bytes()",
        "Lifecycle: collect() → validate() → save()",
        "Auto-registration in CollectorRegistry via __init_subclass__",
        "Politeness pause between API calls (rate_limit_delay)",
    ])

    doc.add_heading("4.2 Collection Lifecycle", level=2)
    add_numbered(doc, [
        "Startup log with parameters",
        "collect() — Raw data retrieval",
        "validate() — Quality check (columns, types)",
        "save() — CSV persistence in data/raw/",
        "End log with summary (rows, duration, status)",
    ])

    doc.add_heading("4.3 Resilience", level=2)
    add_bullets(doc, [
        "If a city fails (weather), the others are still collected",
        "If a series fails (INSEE), the others are still merged",
        "Automatic HTTP retry with exponential backoff",
        "PARTIAL status if some elements failed",
    ])
    doc.add_page_break()

    # =================================================================
    # 5. DATA PROCESSING
    # =================================================================
    doc.add_heading("5. Data Processing", level=1)

    doc.add_heading("5.1 Cleaning (DataCleaner)", level=2)
    doc.add_paragraph(
        "The clean_data.py module applies source-specific cleaning rules. "
        "Cleaning is idempotent."
    )

    doc.add_heading("Weather", level=3)
    add_bullets(doc, [
        "Date conversion to datetime",
        "Duplicate removal (date x city)",
        "Clipping: temperature [-30, 50], precipitation [0, 300], wind [0, 200]",
        "HDD/CDD recalculation (base 18 degrees)",
    ])

    doc.add_heading("DPE", level=3)
    add_bullets(doc, [
        "Processing in chunks of 200,000 rows (memory management)",
        "Deduplication on numero_dpe",
        "Date validation (>= 2021-07-01)",
        "DPE label validation (A-G)",
        "Clipping: surface [5, 1000], consumption [0, 1000], cost [0, 50000]",
        "Heat pump detection: regex on heating and cooling generator type",
        "AC detection: cooling generator field populated",
    ])

    doc.add_heading("5.2 Merging (DatasetMerger)", level=2)
    doc.add_paragraph(
        "DPE (aggregated month x dept) "
        "LEFT JOIN Weather (aggregated month x dept) ON date_id + dept "
        "LEFT JOIN INSEE (month) ON date_id "
        "LEFT JOIN Eurostat (month) ON date_id"
    )
    doc.add_paragraph(
        "Economic indicators (INSEE, Eurostat) are national — "
        "they are broadcast to all departments during the join."
    )

    doc.add_heading("5.3 Outlier Detection", level=2)
    doc.add_paragraph(
        "The outlier_detection.py module provides several methods: "
        "IQR, Z-score, Isolation Forest, LOF. Outliers are "
        "reported but not removed by default (conservative approach)."
    )
    doc.add_page_break()

    # =================================================================
    # 6. FEATURE ENGINEERING
    # =================================================================
    doc.add_heading("6. Feature Engineering", level=1)
    doc.add_paragraph(
        "The FeatureEngineer transforms the dataset into ~90 features. "
        "All temporal operations are performed PER DEPARTMENT "
        "(groupby dept) to avoid cross-department leakage."
    )

    doc.add_heading("6.1 Temporal Lags", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["*_lag_1m", "Previous month value"],
            ["*_lag_3m", "Value from 3 months ago"],
            ["*_lag_6m", "Value from 6 months ago"],
        ]
    )

    doc.add_heading("6.2 Rolling Windows", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["*_rmean_3m / *_rmean_6m", "Rolling mean over 3/6 months"],
            ["*_rstd_3m / *_rstd_6m", "Rolling standard deviation over 3/6 months"],
        ]
    )

    doc.add_heading("6.3 Changes", level=2)
    add_table(doc,
        ["Feature", "Description"],
        [
            ["*_diff_1m", "Absolute difference from previous month"],
            ["*_pct_1m", "Relative change in % (clipped [-200, +500])"],
        ]
    )

    doc.add_heading("6.4 Interaction Features", level=2)
    add_table(doc,
        ["Feature", "Business Hypothesis"],
        [
            ["interact_hdd_confiance", "Cold winter + confidence = heating investment"],
            ["interact_cdd_ipi", "Hot summer + industrial activity = AC installations"],
            ["interact_confiance_bat", "Double proxy of investment intention"],
            ["jours_extremes", "Extreme weather = purchase trigger"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 7. DATABASE SCHEMA
    # =================================================================
    doc.add_heading("7. Database Schema", level=1)
    doc.add_paragraph(
        "Star Schema with SQLAlchemy. "
        "Three supported engines: SQLite (default), PostgreSQL, SQL Server."
    )

    doc.add_heading("7.1 Dimension Tables", level=2)
    add_table(doc,
        ["Table", "Granularity", "Description"],
        [
            ["dim_time", "Month", "Time dimension (YYYYMM, HVAC seasons)"],
            ["dim_geo", "Department", "96 departments with reference city"],
            ["dim_equipment_type", "Type", "HVAC equipment catalog"],
        ]
    )

    doc.add_heading("7.2 Fact Tables", level=2)
    add_table(doc,
        ["Table", "Granularity", "Description"],
        [
            ["fact_hvac_installations", "Month x Dept", "Target variable + weather + permits"],
            ["fact_economic_context", "Month", "National economic indicators"],
            ["raw_dpe", "Individual DPE", "Individual diagnostics"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 8. ML MODELING
    # =================================================================
    doc.add_heading("8. ML Modeling", level=1)

    doc.add_heading("8.1 Temporal Split", level=2)
    doc.add_paragraph("The split respects chronological order (no temporal leakage):")
    add_table(doc,
        ["Split", "Period", "Duration"],
        [
            ["Train", "2021-07 to 2024-06", "36 months"],
            ["Validation", "2024-07 to 2024-12", "6 months"],
            ["Test", "2025-01 to 2025-12+", "12+ months"],
        ]
    )

    doc.add_heading("8.2 Preprocessing", level=2)
    add_numbered(doc, [
        "Separation of features (X) / target (y)",
        "Removal of identifiers (date_id, dept, city_ref)",
        "Median imputation of NaN (SimpleImputer)",
        "Standardization (StandardScaler) for Ridge and LSTM",
    ])

    doc.add_heading("8.3 Models", level=2)

    doc.add_heading("Ridge Regression (robust)", level=3)
    doc.add_paragraph(
        "L2 regularized linear regression. Alpha selected by temporal cross-validation "
        "(TimeSeriesSplit, 3 folds). Negative predictions clipped to 0."
    )

    doc.add_heading("LightGBM (gradient boosting)", level=3)
    doc.add_paragraph(
        "Regularized gradient boosting for small dataset: "
        "max_depth=4, num_leaves=15, min_child_samples=20. "
        "Early stopping on validation (20 rounds)."
    )

    doc.add_heading("Prophet (time series)", level=3)
    doc.add_paragraph(
        "Meta additive model trained per department. "
        "External regressors: temp_mean, hdd_sum, cdd_sum, confiance_menages."
    )

    doc.add_heading("LSTM (exploratory)", level=3)
    doc.add_paragraph(
        "Recurrent neural network (PyTorch). Exploratory only — "
        "the data volume is too small for robust DL."
    )

    doc.add_heading("8.4 Evaluation", level=2)
    add_table(doc,
        ["Metric", "Interpretation"],
        [
            ["RMSE", "Average error in target units"],
            ["MAE", "Mean absolute error"],
            ["MAPE", "Relative error as percentage"],
            ["R2", "Proportion of variance explained (1 = perfect)"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 9. RIDGE AUTOCORRELATION ANALYSIS
    # =================================================================
    doc.add_heading("9. Ridge Auto-correlation Analysis", level=1)

    doc.add_heading("9.1 Observation", level=2)
    doc.add_paragraph(
        "The Ridge model achieves an R2 of 0.998 on the test set. "
        "Coefficient analysis reveals that the dominant features "
        "are the target variable lags (auto-regression)."
    )

    doc.add_heading("9.2 Explanation", level=2)
    doc.add_paragraph(
        "This is not a data leak (lags use past values), "
        "but the model is essentially performing auto-regression. "
        "It predicts month M from month M-1."
    )

    doc.add_heading("9.3 Solution", level=2)
    doc.add_paragraph(
        "The ModelTrainer has an exclude_target_lags=True parameter that "
        "excludes all features derived from the target variable. "
        "A 'ridge_exogenous' model is automatically trained for comparison."
    )
    doc.add_page_break()

    # =================================================================
    # 10. REST API (FastAPI)
    # =================================================================
    doc.add_heading("10. REST API (FastAPI)", level=1)

    doc.add_paragraph(
        "The API exposes ML models via 6 REST endpoints. "
        "It is built with FastAPI and Pydantic v2."
    )

    doc.add_heading("10.1 Architecture", level=2)
    add_table(doc,
        ["File", "Responsibility"],
        [
            ["api/main.py", "Endpoints, CORS middleware, lifespan (model loading)"],
            ["api/models.py", "Pydantic v2 schemas (request/response)"],
            ["api/dependencies.py", "AppState singleton, .pkl model loading"],
        ]
    )

    doc.add_heading("10.2 Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/health", "Health check + version + model status"],
            ["GET", "/predictions", "Stored predictions (with dept/date filters)"],
            ["POST", "/predict", "Custom prediction (manual values)"],
            ["GET", "/data/summary", "Dataset summary (rows, columns, sources)"],
            ["GET", "/model/metrics", "Model evaluation metrics"],
            ["GET", "/departments", "List of 96 departments"],
        ]
    )

    doc.add_heading("10.3 Startup", level=2)
    doc.add_paragraph("uvicorn api.main:app --reload", style="No Spacing")
    doc.add_paragraph("Interactive documentation: http://localhost:8000/docs")
    doc.add_page_break()

    # =================================================================
    # 11. STREAMLIT DASHBOARD
    # =================================================================
    doc.add_heading("11. Streamlit Dashboard", level=1)

    doc.add_paragraph(
        "The Streamlit dashboard provides 6 interactive pages to "
        "explore the data and model results."
    )

    add_table(doc,
        ["Page", "File", "Content"],
        [
            ["Home", "app/pages/home.py", "Project presentation and KPIs"],
            ["Exploration", "app/pages/exploration.py", "Interactive charts, dept/date filters"],
            ["Map", "app/pages/carte.py", "Department choropleth map"],
            ["Models", "app/pages/models.py", "Model comparison, metrics"],
            ["Predictions", "app/pages/predictions.py", "Future predictions, charts"],
            ["Pipeline", "app/pages/pipeline_page.py", "Pipeline execution from the dashboard"],
        ]
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The dashboard uses @st.cache_data(ttl=300) for caching "
        "and a dynamic page loading system."
    )
    doc.add_paragraph("Startup: streamlit run app/app.py", style="No Spacing")
    doc.add_page_break()

    # =================================================================
    # 12. DEPLOYMENT (Docker / Kubernetes)
    # =================================================================
    doc.add_heading("12. Deployment (Docker / Kubernetes)", level=1)

    doc.add_heading("12.1 Docker", level=2)
    doc.add_paragraph("The Dockerfile uses a multi-stage build (3 stages):")
    add_numbered(doc, [
        "base — Python 3.11-slim, system dependencies",
        "dependencies — pip installation (cached layer)",
        "app — Code copy, entrypoint",
    ])

    doc.add_heading("12.2 Docker Compose", level=2)
    add_table(doc,
        ["Service", "Port", "Description"],
        [
            ["api", "8000", "FastAPI (uvicorn)"],
            ["dashboard", "8501", "Streamlit"],
            ["pipeline", "—", "Pipeline execution (profile: tools)"],
            ["postgres", "5432", "PostgreSQL 16 (profile: db, optional)"],
        ]
    )

    doc.add_heading("12.3 Execution Modes", level=2)
    doc.add_paragraph("The docker-entrypoint.sh supports 5 modes:")
    add_table(doc,
        ["Mode", "Command", "Description"],
        [
            ["all", "docker run -e MODE=all ...", "API + Dashboard (default)"],
            ["api", "docker run -e MODE=api ...", "API only"],
            ["dashboard", "docker run -e MODE=dashboard ...", "Dashboard only"],
            ["pipeline", "docker run -e MODE=pipeline ...", "Pipeline execution"],
            ["demo", "docker run -e MODE=demo ...", "Demo data → pipeline → services"],
        ]
    )

    doc.add_heading("12.4 Kubernetes", level=2)
    doc.add_paragraph("6 manifests in the kubernetes/ directory:")
    add_table(doc,
        ["File", "Resource", "Description"],
        [
            ["namespace.yaml", "Namespace", "hvac-market"],
            ["api-deployment.yaml", "Deployment + Service", "2 replicas, probes /health"],
            ["dashboard-deployment.yaml", "Deployment + Service", "1 replica, probes /_stcore/health"],
            ["ingress.yaml", "Ingress", "nginx, /api/* → API, /* → Dashboard"],
            ["pvc.yaml", "PersistentVolumeClaim", "5Gi ReadWriteOnce"],
            ["cronjob-pipeline.yaml", "CronJob", "Monthly pipeline (1st of month, 6am)"],
        ]
    )

    doc.add_heading("12.5 Render (PaaS)", level=2)
    doc.add_paragraph(
        "The render.yaml file defines 2 web services (API + Dashboard) "
        "deployable with one click on Render.com (free tier)."
    )
    doc.add_page_break()

    # =================================================================
    # 13. ORCHESTRATION (Airflow)
    # =================================================================
    doc.add_heading("13. Orchestration (Airflow)", level=1)

    doc.add_paragraph(
        "The Airflow DAG (airflow/dags/hvac_pipeline_dag.py) orchestrates "
        "the complete pipeline with 3 task groups:"
    )

    add_table(doc,
        ["TaskGroup", "Tasks", "Mode"],
        [
            ["collect_group", "5 collectors (weather, insee, eurostat, dpe, sitadel)", "Parallel"],
            ["process_group", "clean → merge → features → eda (6 steps)", "Sequential"],
            ["ml_group", "train + evaluate", "Parallel"],
        ]
    )

    doc.add_paragraph("")
    add_bullets(doc, [
        "Schedule: 0 6 1 * * (1st of month at 6am)",
        "Quality gates between each group (PythonOperator)",
        "SLAs and failure callbacks",
        "Execution_timeout to prevent deadlocks",
    ])
    doc.add_page_break()

    # =================================================================
    # 14. REINFORCEMENT LEARNING
    # =================================================================
    doc.add_heading("14. Reinforcement Learning (demo)", level=1)

    doc.add_paragraph(
        "The module src/models/reinforcement_learning_demo.py implements "
        "a Gymnasium environment simulating HVAC maintenance."
    )

    doc.add_heading("14.1 Environment (HVACMaintenanceEnv)", level=2)
    add_table(doc,
        ["Element", "Description"],
        [
            ["State (6 dim)", "equipment_age, efficiency, temperature, humidity, "
             "nb_failures, cumulative_cost"],
            ["Actions (4)", "0: nothing, 1: light maintenance, 2: full maintenance, "
             "3: replacement"],
            ["Reward", "Efficiency x bonus — action_cost — failure_penalty"],
            ["Episode", "12 months (1 year of maintenance)"],
        ]
    )

    doc.add_heading("14.2 Q-Learning Agent", level=2)
    add_bullets(doc, [
        "Discretization of continuous state into bins",
        "Epsilon-greedy policy (exploration → exploitation)",
        "Comparison with random policy",
        "Visualizations: learning curve, Q-values heatmap",
    ])
    doc.add_page_break()

    # =================================================================
    # 15. UNIT TESTS
    # =================================================================
    doc.add_heading("15. Unit Tests", level=1)

    doc.add_paragraph("The project has 119 unit tests:")
    add_table(doc,
        ["File", "Tests", "Coverage"],
        [
            ["tests/test_config.py", "11+", "GeoConfig, TimeConfig, DatabaseConfig, ModelConfig"],
            ["tests/test_collectors/test_base.py", "12", "CollectorStatus, Config, Result, Registry"],
            ["tests/test_collectors/test_weather.py", "6", "Validation, mocked collection, resilience"],
            ["tests/test_collectors/test_insee.py", "5", "Validation, series structure"],
            ["tests/test_collectors/test_pcloud_sync.py", "~20", "Sync, upload, download, backward compat"],
            ["tests/test_processing/test_clean_data.py", "7", "Weather, INSEE, Eurostat cleaning"],
            ["tests/test_processing/test_feature_engineering.py", "16", "Lags, rolling, interactions"],
            ["tests/test_processing/test_outlier_detection.py", "~20", "IQR, Z-score, IF, LOF"],
            ["tests/test_models/test_robust_scaling.py", "~5", "RobustScaler"],
        ]
    )

    doc.add_paragraph(
        "Tests use shared fixtures (tests/conftest.py) "
        "and HTTP calls are mocked (no network dependency)."
    )
    doc.add_page_break()

    # =================================================================
    # 16. FORMATION MODULES COVERAGE
    # =================================================================
    doc.add_heading("16. Training Modules Coverage", level=1)

    doc.add_paragraph(
        "The project covers the 6 modules of the "
        "Data Science Lead certification (RNCP Master's level, Jedha Bootcamp):"
    )

    add_table(doc,
        ["Module", "Subject", "Implementation in the project"],
        [
            ["M1", "Data Governance",
             "docs/DATA_GOVERNANCE.md — GDPR, AI Act, data lineage, maturity audit"],
            ["M2", "Deployment & Distributed ML",
             "Dockerfile, docker-compose.yml, kubernetes/, render.yaml"],
            ["M3", "Database Architecture",
             "docs/DATABASE_ARCHITECTURE.md — Star schema, OLAP, MongoDB"],
            ["M4", "Data Pipelines",
             "docs/DATA_PIPELINE.md — ETL architecture, monitoring, Airbyte"],
            ["M5", "Automation & Workflow",
             "airflow/dags/hvac_pipeline_dag.py — Complete DAG, TaskGroups"],
            ["M6", "Reinforcement Learning",
             "src/models/reinforcement_learning_demo.py — Gymnasium + Q-Learning"],
        ]
    )

    # Save
    path = "docs/documentation_technique.docx"
    doc.save(path)
    print(f"Technical documentation generated: {path}")
    return path


# ================================================================
# DOCUMENT 2 : User Guide & Configuration
# ================================================================

def generate_user_guide():
    """Generate the user guide with installation, config, and accounts."""
    doc = Document()
    set_doc_styles(doc)
    add_cover_page(
        doc,
        "User Guide",
        "Installation, Configuration, Accounts and Deployment"
    )

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Prerequisites",
        "2. Installation",
        "3. Configuration (.env)",
        "4. Pipeline Commands",
        "5. External Accounts and Services",
        "6. Docker Deployment",
        "7. Kubernetes Deployment",
        "8. Render Deployment (PaaS)",
        "9. Dashboard & API",
        "10. Tests",
        "11. Data Structure",
        "12. Troubleshooting",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # =================================================================
    # 1. PREREQUISITES
    # =================================================================
    doc.add_heading("1. Prerequisites", level=1)
    add_table(doc,
        ["Component", "Version", "Notes"],
        [
            ["Python", "3.10+", "Required"],
            ["Git", "2.x+", "Version control"],
            ["pip", "Latest", "Package manager"],
            ["Docker", "20.x+", "Optional (containerized deployment)"],
            ["Docker Compose", "2.x+", "Optional (multi-services)"],
            ["kubectl", "1.25+", "Optional (Kubernetes deployment)"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 2. INSTALLATION
    # =================================================================
    doc.add_heading("2. Installation", level=1)

    doc.add_heading("2.1 Repository Cloning", level=2)
    doc.add_paragraph(
        "git clone https://github.com/PDUCLOS/Projet-HVAC.git\n"
        "cd Projet-HVAC",
        style="No Spacing"
    )

    doc.add_heading("2.2 Virtual Environment", level=2)
    doc.add_paragraph("Windows:", style="List Bullet")
    doc.add_paragraph(
        "python -m venv venv\nvenv\\Scripts\\activate",
        style="No Spacing"
    )
    doc.add_paragraph("Linux / Mac:", style="List Bullet")
    doc.add_paragraph(
        "python -m venv venv\nsource venv/bin/activate",
        style="No Spacing"
    )

    doc.add_heading("2.3 Dependencies", level=2)
    doc.add_paragraph(
        "# Core (required)\n"
        "pip install -r requirements.txt\n\n"
        "# API (if using FastAPI)\n"
        "pip install -r requirements-api.txt\n\n"
        "# Deep Learning (optional)\n"
        "pip install -r requirements-dl.txt",
        style="No Spacing"
    )

    doc.add_heading("2.4 Quick Setup", level=2)
    doc.add_paragraph("python setup_project.py", style="No Spacing")
    doc.add_paragraph(
        "This script creates the data/ directories, initializes the SQLite database "
        "and verifies that dependencies are installed."
    )
    doc.add_page_break()

    # =================================================================
    # 3. CONFIGURATION
    # =================================================================
    doc.add_heading("3. Configuration (.env)", level=1)

    doc.add_heading("3.1 File Creation", level=2)
    doc.add_paragraph("cp .env.example .env", style="No Spacing")

    doc.add_heading("3.2 Main Variables", level=2)
    add_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["DB_TYPE", "sqlite", "Database type (sqlite, mssql, postgresql)"],
            ["DB_PATH", "data/hvac_market.db", "SQLite file path"],
            ["DATA_START_DATE", "2019-01-01", "Collection start date"],
            ["DATA_END_DATE", "2026-02-28", "Collection end date"],
            ["TARGET_REGION", "FR", "Region code (FR = Metropolitan France, 96 depts)"],
            ["REQUEST_TIMEOUT", "30", "HTTP timeout in seconds"],
            ["MAX_RETRIES", "3", "Maximum number of retries"],
            ["LOG_LEVEL", "INFO", "Logging level"],
        ]
    )

    doc.add_heading("3.3 Database Configuration", level=2)
    doc.add_paragraph(
        "# SQLite (default — no installation needed)\n"
        "DB_TYPE=sqlite\n"
        "DB_PATH=data/hvac_market.db\n\n"
        "# PostgreSQL (production/cloud)\n"
        "DB_TYPE=postgresql\n"
        "DB_HOST=localhost\n"
        "DB_PORT=5432\n"
        "DB_NAME=hvac\n"
        "DB_USER=hvac_user\n"
        "DB_PASSWORD=your_password\n"
        "ALLOW_NON_LOCAL=true",
        style="No Spacing"
    )
    doc.add_page_break()

    # =================================================================
    # 4. PIPELINE COMMANDS
    # =================================================================
    doc.add_heading("4. Pipeline Commands", level=1)

    add_table(doc,
        ["Command", "Description", "Duration"],
        [
            ["python -m src.pipeline list", "List collectors", "< 1s"],
            ["python -m src.pipeline init_db", "Initialize the database", "< 5s"],
            ["python -m src.pipeline collect", "Collect all sources", "30-90 min"],
            ["python -m src.pipeline collect --sources weather,insee",
             "Collect specific sources", "< 5 min"],
            ["python -m src.pipeline clean", "Clean raw data", "1-3 min"],
            ["python -m src.pipeline merge", "Merge into ML dataset", "< 30s"],
            ["python -m src.pipeline features", "Feature engineering", "< 10s"],
            ["python -m src.pipeline process", "clean + merge + features", "2-5 min"],
            ["python -m src.pipeline eda", "Exploratory analysis", "1-2 min"],
            ["python -m src.pipeline train", "Train ML models", "2-5 min"],
            ["python -m src.pipeline evaluate", "Evaluate models", "3-5 min"],
            ["python -m src.pipeline all", "Complete pipeline", "40-90 min"],
        ]
    )

    doc.add_heading("Makefile Shortcuts", level=2)
    add_table(doc,
        ["Command", "Action"],
        [
            ["make install", "Create venv + install dependencies"],
            ["make collect", "Collect all sources"],
            ["make process", "Process data"],
            ["make train", "Train models"],
            ["make all", "Complete pipeline"],
            ["make test", "Run all 119 tests"],
            ["make api", "Start the FastAPI API"],
            ["make dashboard", "Start the Streamlit dashboard"],
            ["make docker-build", "Build the Docker image"],
            ["make docker-run", "Run via Docker"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 5. ACCOUNTS AND EXTERNAL SERVICES
    # =================================================================
    doc.add_heading("5. External Accounts and Services", level=1)

    doc.add_paragraph(
        "Here are the accounts and configurations required depending on your usage:"
    )

    doc.add_heading("5.1 No account required (local usage)", level=2)
    doc.add_paragraph(
        "To use the project locally, NO account is needed. "
        "All data sources are Open Data without API keys."
    )
    add_bullets(doc, [
        "DPE ADEME: free access (data.ademe.fr)",
        "Open-Meteo: free access (archive-api.open-meteo.com)",
        "INSEE BDM: free access (bdm.insee.fr/sdmx)",
        "Eurostat: free access (ec.europa.eu/eurostat)",
    ])

    doc.add_heading("5.2 GitHub (version control + CI/CD)", level=2)
    add_table(doc,
        ["Element", "Action"],
        [
            ["Account", "Create an account on github.com (free)"],
            ["Repository", "Fork or clone PDUCLOS/Projet-HVAC"],
            ["Visibility", "Set the repo to PUBLIC for the portfolio"],
            ["GitHub Pages", "Optional: enable to host documentation"],
        ]
    )

    doc.add_heading("5.3 Render.com (PaaS deployment)", level=2)
    add_table(doc,
        ["Element", "Action"],
        [
            ["Account", "Create an account on render.com (free)"],
            ["Connection", "Link with your GitHub account"],
            ["Deployment", "New > Blueprint > select the repo"],
            ["Blueprint", "render.yaml configures 2 services automatically"],
            ["Tier", "Free tier (sufficient for demo/portfolio)"],
            ["Variables", "Add DB_TYPE=sqlite in Environment"],
        ]
    )

    doc.add_heading("5.4 Docker Hub (Docker images)", level=2)
    add_table(doc,
        ["Element", "Action"],
        [
            ["Account", "Create an account on hub.docker.com (free)"],
            ["Login", "docker login"],
            ["Build", "docker build -t your-user/hvac-market ."],
            ["Push", "docker push your-user/hvac-market"],
        ]
    )
    doc.add_paragraph(
        "Note: Docker Hub is only needed if you want to distribute "
        "the image. For local usage, 'docker build' is sufficient."
    )

    doc.add_heading("5.5 Cloud Kubernetes (optional)", level=2)
    doc.add_paragraph(
        "To deploy on a Kubernetes cluster:"
    )
    add_table(doc,
        ["Provider", "Service", "Free tier"],
        [
            ["Google Cloud", "GKE Autopilot", "$300 credits (90 days)"],
            ["AWS", "EKS", "No free tier for EKS"],
            ["Azure", "AKS", "Free tier control plane"],
            ["OVHcloud", "Managed Kubernetes", "Free (pay-as-you-go nodes)"],
        ]
    )
    doc.add_paragraph(
        "For the portfolio, Render.com (free) is recommended. "
        "Kubernetes is documented to demonstrate the skill."
    )

    doc.add_heading("5.6 pCloud (data storage)", level=2)
    add_table(doc,
        ["Element", "Action"],
        [
            ["Account", "Create an account on pcloud.com (free, 10 GB)"],
            ["Configuration", "Add PCLOUD_USER and PCLOUD_PASS in .env"],
            ["Usage", "python -m src.pipeline sync — synchronizes data/"],
            ["Optional", "Only if you work on multiple machines"],
        ]
    )

    doc.add_heading("5.7 Streamlit Cloud (online dashboard)", level=2)
    add_table(doc,
        ["Element", "Action"],
        [
            ["Account", "Create an account on streamlit.io (free)"],
            ["Connection", "Link with your GitHub account"],
            ["Deployment", "New app > select repo > app/app.py"],
            ["Data", "Use data/sample/ for demo mode"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 6. DOCKER DEPLOYMENT
    # =================================================================
    doc.add_heading("6. Docker Deployment", level=1)

    doc.add_heading("6.1 Build", level=2)
    doc.add_paragraph("docker build -t hvac-market .", style="No Spacing")

    doc.add_heading("6.2 Run (full mode)", level=2)
    doc.add_paragraph(
        "docker compose up -d\n"
        "# API: http://localhost:8000\n"
        "# Dashboard: http://localhost:8501",
        style="No Spacing"
    )

    doc.add_heading("6.3 Run (demo mode)", level=2)
    doc.add_paragraph(
        "docker run -e MODE=demo -p 8000:8000 -p 8501:8501 hvac-market",
        style="No Spacing"
    )
    doc.add_paragraph(
        "Demo mode generates synthetic data, runs the pipeline, "
        "then launches the API and dashboard automatically."
    )
    doc.add_page_break()

    # =================================================================
    # 7. KUBERNETES DEPLOYMENT
    # =================================================================
    doc.add_heading("7. Kubernetes Deployment", level=1)

    doc.add_paragraph("Apply the manifests in order:")
    doc.add_paragraph(
        "kubectl apply -f kubernetes/namespace.yaml\n"
        "kubectl apply -f kubernetes/pvc.yaml\n"
        "kubectl apply -f kubernetes/api-deployment.yaml\n"
        "kubectl apply -f kubernetes/dashboard-deployment.yaml\n"
        "kubectl apply -f kubernetes/ingress.yaml\n"
        "kubectl apply -f kubernetes/cronjob-pipeline.yaml",
        style="No Spacing"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Verify the deployment:\n"
        "kubectl get pods -n hvac-market\n"
        "kubectl get svc -n hvac-market",
        style="No Spacing"
    )
    doc.add_page_break()

    # =================================================================
    # 8. RENDER DEPLOYMENT
    # =================================================================
    doc.add_heading("8. Render Deployment (PaaS)", level=1)

    doc.add_heading("Steps", level=2)
    add_numbered(doc, [
        "Create an account on render.com and link GitHub",
        "Click 'New > Blueprint'",
        "Select the repository PDUCLOS/Projet-HVAC",
        "Render automatically detects render.yaml",
        "2 services are created: hvac-api (FastAPI) and hvac-dashboard (Streamlit)",
        "Add environment variables if needed (DB_TYPE=sqlite)",
        "Click Deploy — URLs are generated automatically",
    ])
    doc.add_page_break()

    # =================================================================
    # 9. DASHBOARD & API
    # =================================================================
    doc.add_heading("9. Dashboard & API", level=1)

    doc.add_heading("9.1 Local Launch", level=2)
    doc.add_paragraph(
        "# API (terminal 1)\n"
        "uvicorn api.main:app --reload --port 8000\n\n"
        "# Dashboard (terminal 2)\n"
        "streamlit run app/app.py --server.port 8501",
        style="No Spacing"
    )

    doc.add_heading("9.2 API Documentation", level=2)
    add_bullets(doc, [
        "Swagger UI : http://localhost:8000/docs",
        "ReDoc : http://localhost:8000/redoc",
        "OpenAPI JSON : http://localhost:8000/openapi.json",
    ])
    doc.add_page_break()

    # =================================================================
    # 10. TESTS
    # =================================================================
    doc.add_heading("10. Tests", level=1)

    doc.add_paragraph("python -m pytest tests/ -v", style="No Spacing")
    doc.add_paragraph("")
    doc.add_paragraph(
        "# With coverage\n"
        "pip install pytest-cov\n"
        "python -m pytest tests/ --cov=src --cov=config --cov-report=term-missing",
        style="No Spacing"
    )
    doc.add_page_break()

    # =================================================================
    # 11. DATA STRUCTURE
    # =================================================================
    doc.add_heading("11. Data Structure", level=1)

    add_table(doc,
        ["Directory", "Content"],
        [
            ["data/raw/weather/", "weather_france.csv — daily weather data"],
            ["data/raw/insee/", "indicateurs_economiques.csv — INSEE series"],
            ["data/raw/eurostat/", "ipi_hvac_france.csv — production indices"],
            ["data/raw/dpe/", "dpe_france_all.csv — raw DPE data"],
            ["data/processed/", "Cleaned data"],
            ["data/features/", "hvac_ml_dataset.csv and hvac_features_dataset.csv"],
            ["data/models/", "Models .pkl, metrics, reports"],
            ["data/models/figures/", "Visualization charts"],
            ["data/analysis/", "EDA results and correlations"],
            ["data/sample/", "Sample data for demo (200 rows)"],
        ]
    )
    doc.add_page_break()

    # =================================================================
    # 12. TROUBLESHOOTING
    # =================================================================
    doc.add_heading("12. Troubleshooting", level=1)

    problems = [
        (
            "ModuleNotFoundError",
            "pip install -r requirements.txt\n"
            "Verify that the virtual environment is activated."
        ),
        (
            "DPE collection fails or is slow",
            "The ADEME API may be overloaded. Re-run the command. "
            "The collector is idempotent."
        ),
        (
            "FileNotFoundError: Dataset not found",
            "Steps must be executed in order: "
            "collect → process → train. "
            "Run: python -m src.pipeline process"
        ),
        (
            "Ridge R2 > 0.99",
            "Normal with lag features (auto-regression). "
            "See the 'ridge_exogenous' model for performance "
            "without auto-correlation."
        ),
        (
            "Docker build fails",
            "Verify that Docker is installed and running. "
            "Check disk space (>5 GB recommended)."
        ),
        (
            "API does not start",
            "Verify that .pkl models exist in data/models/. "
            "Otherwise, run: python -m src.pipeline train"
        ),
    ]

    for problem, solution in problems:
        doc.add_heading(problem, level=3)
        doc.add_paragraph(solution)

    doc.add_heading("Contact", level=2)
    doc.add_paragraph(
        "Author: Patrice DUCLOS — Senior Data Analyst\n"
        "GitHub Repository: https://github.com/PDUCLOS/Projet-HVAC\n"
        "Program: Data Science Lead — Jedha Bootcamp"
    )

    # Save
    path = "docs/guide_utilisation.docx"
    doc.save(path)
    print(f"User guide generated: {path}")
    return path


# ================================================================
# MAIN
# ================================================================

if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    generate_technical_doc()
    generate_user_guide()
    print("\nDocumentation generated successfully in docs/")
